#!/usr/bin/env python3
"""
PDFlexity OCR Worker — PaddleOCR-powered PDF OCR pipeline.

This script is spawned by the Go pdf-engine as a subprocess.
Communication protocol:
  - Args: --input, --output-dir, --languages, --dpi, --mode, --page, --export-format, --export-output
  - Stdout: one JSON object per line (streaming events)
  - Stderr: debug logs (never parsed by caller)

Modes:
  full         — Full OCR pipeline: render all pages, run OCR, stream results
  render-page  — Render a single page as base64 image
  export       — Export OCR results to PDF/DOCX/JSON (reads OCR data from stdin)
"""

import argparse
import json
import io
import time
import base64
import os
import sys
import uuid
import logging
from typing import Dict, Any, List

# Force stdout to be UTF-8 to avoid UnicodeEncodeError on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

logging.basicConfig(stream=sys.stderr, level=logging.INFO, format="[ocr-worker] %(message)s")
logger = logging.getLogger(__name__)


def emit(event: dict):
    """Write a JSON event to stdout (one line, flushed immediately)."""
    sys.stdout.write(json.dumps(event, ensure_ascii=False) + "\n")
    sys.stdout.flush()


def emit_progress(status: str, current_page: int = 0, total_pages: int = 0):
    emit({
        "type": "progress",
        "status": status,
        "currentPage": current_page,
        "totalPages": total_pages,
    })


def emit_error(msg: str):
    emit({"type": "error", "error": msg})


def classify_block_type(text: str, bbox: dict, page_width: float, page_height: float, font_size: float) -> str:
    """Heuristic classification of a text block's type."""
    text_stripped = text.strip()
    
    # Heading detection: large font, short text, near top
    if font_size >= 18 and len(text_stripped) < 120:
        return "heading"
    if font_size >= 14 and len(text_stripped) < 80 and bbox["y"] < page_height * 0.15:
        return "heading"
    
    # Footer/header detection
    if bbox["y"] > page_height * 0.92:
        return "footer"
    if bbox["y"] < page_height * 0.06 and font_size < 12:
        return "header"
    
    # List item detection
    if text_stripped and (text_stripped[0] in "•·-–—◦▪►" or 
        (len(text_stripped) > 2 and text_stripped[0].isdigit() and text_stripped[1] in ".)" )):
        return "list-item"
    
    # Caption: small text near bottom or very short
    if font_size <= 10 and len(text_stripped) < 60:
        return "caption"
    
    return "paragraph"


def estimate_font_size(bbox_height: float, text: str) -> float:
    """Estimate font size from bounding box height."""
    lines = max(1, text.count("\n") + 1)
    size = (bbox_height / lines) * 0.75  # Approximate pt from pixel height
    return max(8, min(72, round(size, 1)))


def estimate_alignment(bbox: dict, page_width: float) -> str:
    """Estimate text alignment from position."""
    center_x = bbox["x"] + bbox["width"] / 2
    if abs(center_x - page_width / 2) < page_width * 0.05:
        return "center"
    if bbox["x"] > page_width * 0.6:
        return "right"
    return "left"


def run_full_ocr(input_path: str, output_dir: str, languages: str, dpi: int):
    """Full OCR pipeline: render pages → OCR → stream results."""
    import fitz  # PyMuPDF
    
    logger.info(f"Opening PDF: {input_path}")
    emit_progress("uploading", 0, 0)
    
    try:
        doc = fitz.open(input_path)
    except Exception as e:
        emit_error(f"Failed to open PDF: {e}")
        return
    
    total_pages = len(doc)
    logger.info(f"PDF has {total_pages} pages")
    emit_progress("rendering", 0, total_pages)
    
    # Initialize PaddleOCR
    try:
        from paddleocr import PaddleOCR
        
        lang_map = {
            "en": "en", "hi": "hi", "ja": "japan", "ar": "ar",
            "zh": "ch", "ko": "korean", "fr": "french", "de": "german",
            "es": "es", "pt": "pt", "ru": "ru", "it": "it",
        }
        
        lang_list = [l.strip() for l in languages.split(",")]
        paddle_lang = lang_map.get(lang_list[0], "en") if lang_list else "en"
        logger.info(f"Initializing PaddleOCR with lang={paddle_lang}")
        ocr_engine = PaddleOCR(
            use_textline_orientation=True,
            lang=paddle_lang,
        )
    except Exception as e:
        emit_error(f"Failed to initialize PaddleOCR: {e}")
        doc.close()
        return
    
    emit_progress("detecting-layout", 0, total_pages)
    
    all_languages = set()
    total_confidence = 0.0
    start_time = time.time()
    
    for page_idx in range(total_pages):
        page_start = time.time()
        page_num = page_idx + 1
        
        emit_progress("running-ocr", page_num, total_pages)
        
        page = doc[page_idx]
        page_width = page.rect.width
        page_height = page.rect.height
        
        # Render page
        zoom = 2.0  # Use an integer multiplier to avoid odd dimensions which might crash the CNN
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")
        
        # Emit page image for UI
        emit({
            "type": "page-image",
            "pageImage": {
                "page": page_num,
                "imageBase64": base64.b64encode(img_data).decode("utf-8"),
                "width": page_width,
                "height": page_height
            }
        })
        
        # Run PaddleOCR
        try:
            import numpy as np
            # Make the array writable by copying
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n).copy()
            
            result = ocr_engine.predict(img_array)
            logger.info("Predict finished successfully")
        except Exception as e:
            logger.error(f"OCR failed on page {page_num}: {e}")
            # Emit empty page result
            emit({
                "type": "page-result",
                "pageResult": {
                    "page": page_num,
                    "width": page_width,
                    "height": page_height,
                    "textBlocks": [],
                    "tables": [],
                    "images": [],
                    "language": lang_list[0] if lang_list else "en",
                    "avgConfidence": 0,
                    "processingTimeMs": int((time.time() - page_start) * 1000),
                }
            })
            continue
        
        # Parse OCR results into structured text blocks
        text_blocks = []
        page_confidences = []
        
        if result and len(result) > 0:
            res_dict = result[0]
            dt_polys = res_dict.get('dt_polys', [])
            rec_texts = res_dict.get('rec_texts', [])
            rec_scores = res_dict.get('rec_scores', [])
            try:
                for idx, poly in enumerate(dt_polys):
                    if idx >= len(rec_texts) or idx >= len(rec_scores):
                        continue
                    
                    text = rec_texts[idx]
                    confidence = float(rec_scores[idx])
                    
                    if not text or not str(text).strip():
                        continue
                    
                    # Convert PaddleOCR coordinates (pixel space) to PDF coordinates
                    xs = [float(c[0]) for c in poly]
                    ys = [float(c[1]) for c in poly]
                    
                    # Convert from rendered pixel coords to PDF point coords
                    bbox = {
                        "x": min(xs) / zoom,
                        "y": min(ys) / zoom,
                        "width": (max(xs) - min(xs)) / zoom,
                        "height": (max(ys) - min(ys)) / zoom
                    }
                    
                    font_size = bbox["height"] * 0.8
                    
                    alignment = "left"
                    
                    # Basic classification
                    type_str = "paragraph"
                    text_str = str(text)
                    if len(text_str) > 0 and text_str.isupper() and len(text_str.split()) < 10:
                        type_str = "heading"
                    elif text_str.strip().startswith(("-", "•", "1.", "2.")):
                        type_str = "list"
                        
                    block = {
                        "id": str(uuid.uuid4()),
                        "type": type_str,
                        "text": text_str,
                        "confidence": confidence,
                        "bbox": bbox,
                        "fontFamily": "Inter, sans-serif",
                        "fontSize": round(font_size, 1),
                        "fontWeight": 600 if type_str == "heading" else 400,
                        "fontStyle": "normal",
                        "alignment": alignment,
                        "lineHeight": round(font_size * 1.4, 1),
                        "color": "#000000",
                    }
                    
                    text_blocks.append(block)
                    page_confidences.append(confidence)
            except Exception as e:
                logger.error(f"Error parsing OCR results on page {page_num}: {e}")
        
        avg_confidence = sum(page_confidences) / len(page_confidences) if page_confidences else 0
        total_confidence += avg_confidence
        
        detected_lang = lang_list[0] if lang_list else "en"
        all_languages.add(detected_lang)
        
        processing_time = int((time.time() - page_start) * 1000)
        
        # Emit page result
        emit({
            "type": "page-result",
            "pageResult": {
                "page": page_num,
                "width": page_width,
                "height": page_height,
                "textBlocks": text_blocks,
                "tables": [],  # Table detection can be added with PP-StructureV3
                "images": [],
                "language": detected_lang,
                "avgConfidence": round(avg_confidence, 4),
                "processingTimeMs": processing_time,
            }
        })
    
    doc.close()
    
    overall_confidence = total_confidence / total_pages if total_pages > 0 else 0
    total_time = int((time.time() - start_time) * 1000)
    
    emit({
        "type": "complete",
        "overallConfidence": round(overall_confidence, 4),
        "detectedLanguages": list(all_languages),
        "totalProcessingTimeMs": total_time,
        "data": {
            "totalPages": total_pages,
            "overallConfidence": round(overall_confidence, 4),
            "detectedLanguages": list(all_languages),
        }
    })


def run_render_page(input_path: str, page: int, dpi: int):
    """Render a single PDF page as a base64 PNG image."""
    import fitz
    
    try:
        doc = fitz.open(input_path)
        if page < 1 or page > len(doc):
            emit_error(f"Page {page} out of range (1-{len(doc)})")
            return
        
        pg = doc[page - 1]
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = pg.get_pixmap(matrix=mat, alpha=False)
        
        img_data = pix.tobytes("png")
        b64 = base64.b64encode(img_data).decode("utf-8")
        
        # Output single JSON line
        result = {
            "page": page,
            "imageBase64": b64,
            "width": pix.width,
            "height": pix.height,
        }
        sys.stdout.write(json.dumps(result) + "\n")
        sys.stdout.flush()
        doc.close()
    except Exception as e:
        emit_error(f"Failed to render page: {e}")


def run_export(input_path: str, output_dir: str, export_format: str, export_output: str):
    """Export OCR results to the requested format."""
    
    # Read OCR data from stdin
    ocr_data_line = sys.stdin.readline().strip()
    edits_line = sys.stdin.readline().strip() if True else "{}"
    
    try:
        ocr_data = json.loads(ocr_data_line) if ocr_data_line else {}
    except:
        ocr_data = {}
    
    try:
        edits = json.loads(edits_line) if edits_line else {}
    except:
        edits = {}
    
    if export_format == "json":
        # Export raw OCR structure as JSON
        with open(export_output, "w", encoding="utf-8") as f:
            json.dump(ocr_data, f, indent=2, ensure_ascii=False)
        emit({"type": "complete", "data": {"outputPath": export_output}})
        
    elif export_format == "docx":
        export_docx(ocr_data, edits, export_output)
        
    elif export_format in ("editable-pdf", "searchable-pdf"):
        export_pdf(input_path, ocr_data, edits, export_output, export_format)
        
    else:
        emit_error(f"Unsupported export format: {export_format}")


def export_docx(ocr_data: dict, edits: dict, output_path: str):
    """Export OCR results as a DOCX document."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Process page results
        pages = ocr_data if isinstance(ocr_data, list) else [ocr_data]
        
        for page_data in pages:
            if not isinstance(page_data, dict):
                continue
                
            text_blocks = page_data.get("textBlocks", [])
            
            # Sort blocks by Y position, then X position (reading order)
            text_blocks.sort(key=lambda b: (b.get("bbox", {}).get("y", 0), b.get("bbox", {}).get("x", 0)))
            
            for block in text_blocks:
                text = block.get("text", "").strip()
                if not text:
                    continue
                
                # Apply user edits if any
                block_id = block.get("id", "")
                if block_id in edits:
                    text = edits[block_id].get("text", text)
                
                block_type = block.get("type", "paragraph")
                font_size = block.get("fontSize", 12)
                alignment = block.get("alignment", "left")
                
                if block_type == "heading":
                    heading_level = 1 if font_size >= 20 else 2 if font_size >= 16 else 3
                    para = doc.add_heading(text, level=heading_level)
                else:
                    para = doc.add_paragraph(text)
                    
                    # Set font size
                    for run in para.runs:
                        run.font.size = Pt(min(font_size, 36))
                
                # Set alignment
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        doc.save(output_path)
        emit({"type": "complete", "data": {"outputPath": output_path}})
        
    except Exception as e:
        emit_error(f"DOCX export failed: {e}")


def export_pdf(input_path: str, ocr_data: dict, edits: dict, output_path: str, mode: str):
    """Export as searchable or editable PDF by overlaying text."""
    try:
        import fitz
        
        doc = fitz.open(input_path)
        
        pages = ocr_data if isinstance(ocr_data, list) else [ocr_data]
        
        for page_data in pages:
            if not isinstance(page_data, dict):
                continue
            
            page_num = page_data.get("page", 1) - 1
            if page_num < 0 or page_num >= len(doc):
                continue
            
            page = doc[page_num]
            text_blocks = page_data.get("textBlocks", [])
            
            for block in text_blocks:
                text = block.get("text", "").strip()
                if not text:
                    continue
                
                block_id = block.get("id", "")
                if block_id in edits:
                    text = edits[block_id].get("text", text)
                
                bbox = block.get("bbox", {})
                x = bbox.get("x", 0)
                y = bbox.get("y", 0)
                w = bbox.get("width", 100)
                h = bbox.get("height", 20)
                font_size = block.get("fontSize", 11)
                
                rect = fitz.Rect(x, y, x + w, y + h)
                
                if mode == "searchable-pdf":
                    # Invisible text layer
                    page.insert_text(
                        fitz.Point(x, y + h),
                        text,
                        fontsize=font_size,
                        render_mode=3,  # Invisible
                    )
                else:
                    # Editable: add visible text
                    # First white-out the area
                    page.draw_rect(rect, color=None, fill=(1, 1, 1))
                    page.insert_textbox(
                        rect,
                        text,
                        fontsize=min(font_size, 36),
                        align=0,
                    )
        
        doc.save(output_path)
        doc.close()
        emit({"type": "complete", "data": {"outputPath": output_path}})
        
    except Exception as e:
        emit_error(f"PDF export failed: {e}")


def main():
    parser = argparse.ArgumentParser(description="PDFlexity OCR Worker")
    parser.add_argument("--input", required=True, help="Input PDF path")
    parser.add_argument("--output-dir", default=".", help="Output directory")
    parser.add_argument("--languages", default="en", help="Comma-separated language codes")
    parser.add_argument("--dpi", type=int, default=300, help="Render DPI")
    parser.add_argument("--mode", default="full", choices=["full", "render-page", "export"])
    parser.add_argument("--page", type=int, default=1, help="Page number for render-page mode")
    parser.add_argument("--export-format", default="searchable-pdf", help="Export format")
    parser.add_argument("--export-output", default="", help="Export output path")
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        emit_error(f"Input file not found: {args.input}")
        return
    
    if args.mode == "full":
        run_full_ocr(args.input, args.output_dir, args.languages, args.dpi)
    elif args.mode == "render-page":
        run_render_page(args.input, args.page, args.dpi)
    elif args.mode == "export":
        run_export(args.input, args.output_dir, args.export_format, args.export_output)


if __name__ == "__main__":
    main()
